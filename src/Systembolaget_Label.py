import re, sys, urllib.request
from datetime import date
from docx import Document

# Class to store data about the wine
class wine:
    def __init__(
        self,
        name,
        name2,
        location,
        price,
        productnr,
        alcohol_procentage,
        suger_amount,
        taste_and_usage,
    ):
        self.name = name
        self.name2 = name2
        self.location = location
        self.price = price
        self.productnr = productnr
        self.alcohol_procentage = alcohol_procentage
        self.suger_amount = suger_amount
        self.taste_and_usage = taste_and_usage


# Input 1 string
# Output list o
def load_url_file(file_to_load):
    urls = []
    try:
        urls = open(file_to_load, "r", encoding="utf-8")  # Opens the file with the URLs
    except:
        print(
            "Missing input txt file. Creating that now as input.txt. PLease enter the URLS, row separated"
        )

        new_file = open("input.txt", "w")
        new_file.write("\n")
        new_file.close()

        sys.exit()

    file_content = urls.readlines()
    for row in range(len(file_content)):
        file_content[row] = file_content[row].strip("\n")

    # file_content = infil.read()
    return file_content


# Input list with user input
# Output None, removes invalid input and prints what input was wrong and then returns the fixed list
def input_validation(list_of_pages):

    # Regular expression to practice, easier way is probably too just check if the first X characters of the string match one of the valid url versions.
    pattern_regex_website = re.compile(
        r"^(https:\/\/www.systembolaget.se/produkt/|http:\/\/www.systembolaget.se/produkt/|https:\/\/www.systembolaget.se/\d*)"
    )
    # pattern_regex_generated = re.compile(r'^https://www\.systembolaget\.se/produkt/[a-zA-Z]+/[a-zA-Z]+-[0-9]+/$')

    # check for number inputs
    pattern_regex_nr = re.compile(r"^(\d*)")
    for iterator in range(len(list_of_pages)):
        result = pattern_regex_nr.findall(list_of_pages[iterator])
        if result[0] != "":
            list_of_pages[iterator] = "https://www.systembolaget.se/" + str(
                list_of_pages[iterator]
            )

    # check website inputs including the constructed ones in the previous loop
    validated_pages = []
    index = 0
    for page in list_of_pages:
        result = pattern_regex_website.findall(page)
        if result != []:
            validated_pages.append(page)
        else:
            print("Link on row " + str(index + 1) + " is broken.")

        index += 1

    return validated_pages


# Input URL and index which is used to name it
# Output
def dataurl_data_fetcher(url, index):

    req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
    temp = urllib.request.urlopen(req).read()
    temp = temp.decode(
        "utf-8"
    )  # Website contains ÅÄÖ so we need to ensure proper handling and use UTF-8 decoding since the temp comes in as a datastream of 0&1

    return temp  # site_data


"""
    webcontent = urllib.request.urlopen(url)
    with open( index, "w") as fil:
        for row in webcontent:
            utf8line = row.decode('utf8')
            fil.write(utf8line)
"""

# Input string version of the site
# Output instance of class wine with the data filed.
def extract_drink_data(page_html):

    try:
        # Gets title name
        # wine_name = soup.find("div", class_="react-no-print")
        pattern_wine_name = re.compile(
            r"\"productNameBold\":\"[a-zA-Z\s\S]*\"\,\"productNameThin"
        )  # (r'(\"productNameBold\":\")([a-zA-Z\s])*(",")')
        wine_name = pattern_wine_name.findall(str(page_html))
        wine_name = wine_name[0][19:-18]
        # print(wine_name)
    except IndexError:
        wine_name = "Failed to fetch wine name. Regex 1 failure. "
    except Exception as e:
        wine_name = f"Failed to fetch wine name. {e} "

    try:
        # Gets sub name
        pattern_wine_name2 = re.compile(
            r"\"metaData\"\:\{\"metaTitle\"\:\"[a-zA-Z\s\S]*\|"
        )
        wine_name2 = pattern_wine_name2.findall(str(page_html))
        wine_name2 = wine_name2[0][26 + len(wine_name) : -2]
        # print(wine_name2)
    except IndexError:
        wine_name2 = "Failed to fetch wine sub name. Regex 2 failure. "
    except Exception as e:
        wine_name2 = f"Failed to fetch wine sub name. {e} "

    try:
        # Gets the country
        pattern_wine_location1 = re.compile(
            r"\"country\"\:\"[a-zA-Z\s\S]*\",\"originLevel2"
        )
        wine_location1 = pattern_wine_location1.findall(str(page_html))
        wine_location1 = wine_location1[0][11:-15]
        # print(wine_location1)
    except IndexError:
        wine_location1 = "Failed to fetch wine location 1. Regex 3 failure. "
    except Exception as e:
        wine_location1 = f"Failed to fetch wine location 1. {e} "

    try:
        # Gets the region in the country
        pattern_wine_location2 = re.compile(
            r"\"originLevel1\"\:\"[a-zA-Z\s\S]*\",\"originLevel3"
        )
        wine_location2 = pattern_wine_location2.findall(str(page_html))
        wine_location2 = wine_location2[0][16:-15]
        # print(wine_location2)
    except IndexError:
        wine_location2 = "Failed to fetch wine location 2. Regex 4 failure. "
    except Exception as e:
        wine_location2 = f"Failed to fetch wine location 2. {e} "

    # Combine the locations
    wine_location = wine_location1 + " " + wine_location2
    # print(wine_location)

    try:
        # Gets the wine price
        # pattern_wine_price = re.compile(r'priceInclVat":\d\d.\d\d,\"')
        pattern_wine_price = re.compile(r'priceInclVat":\d*.\d*,\"')
        wine_price = pattern_wine_price.findall(str(page_html))
        wine_price = wine_price[0][14:-2]
        # print(wine_price)
    except IndexError:
        wine_price = "Failed to fetch wine price. Regex 5 failure. "
    except Exception as e:
        wine_price = f"Failed to fetch wine price. {e} "

    try:
        # Gets the systembolaget productnr
        pattern_wine_productnr = re.compile(r'\"productNumberShort\":"\d*\"\,\"volume')
        wine_productnr = pattern_wine_productnr.findall(str(page_html))
        wine_productnr = wine_productnr[0][22:-9]
        # print(wine_productnr)
    except IndexError:
        wine_productnr = "Failed to fetch wine product number. Regex 6 failure. "
    except Exception as e:
        wine_productnr = f"Failed to fetch wine product number. {e} "

    try:
        # Gets the alcohole procentage
        # pattern_wine_alc_procentage = re.compile(r'\"alcoholPercentage\"\:\d\d.\d\d\,\"tastingDate\"')
        pattern_wine_alc_procentage = re.compile(
            r"\"alcoholPercentage\"\:\d*.\d*\,\"tastingDate\""
        )
        wine_alc_procentage = pattern_wine_alc_procentage.findall(str(page_html))
        wine_alc_procentage = wine_alc_procentage[0][20:-14]
        # print(wine_alc_procentage)
    except IndexError:
        wine_alc_procentage = (
            "Failed to fetch wine alcohole procentage. Regex 7 failure. "
        )
    except Exception as e:
        wine_alc_procentage = f"Failed to fetch wine alcohole procentage. {e} "

    try:
        # Get the suger content of the drink
        # pattern_wine_suger_content = re.compile(r'\"sugarContent\"\:\"\d*\"\,\"tasteSymbolsList\"\:')
        # pattern_wine_suger_content = re.compile(r'\"sugarContent\"\:\"\s\d*\"\,\"tasteSymbolsList\"\:')
        # pattern_wine_suger_content = re.compile(r'\"sugarContent\"\:\"\S\d*\"\,\"additives\"\:|\"sugarContent\"\:\"\d*\"\,\"additives\"\:')
        pattern_wine_suger_content = re.compile(
            r"\"sugarContentGramPer100ml\"\:\"\s*\d*\S*\d*\"\,\"additives\"\:"
        )
        wine_suger_content = pattern_wine_suger_content.findall(str(page_html))
        wine_suger_content = wine_suger_content[0][28:-14]
        # print(wine_suger_content)
    except IndexError:
        wine_suger_content = "Failed to fetch wine sugar content. Regex 8 failure. "
    except Exception as e:
        wine_suger_content = f"Failed to fetch wine sugar content. {e} "

    try:
        # Gets the taste and usage recommendations
        pattern_wine_taste = re.compile(r"\"taste\"\:\"[a-zA-Z\s\S]*\"\,\"aroma\"")
        wine_taste = pattern_wine_taste.findall(str(page_html))
        wine_taste = wine_taste[0][9:-9]
        # print(wine_taste_and_usage)
    except IndexError:
        wine_taste = "Failed to fetch wine taste. Regex 9 failure. "
    except Exception as e:
        wine_taste = f"Failed to fetch wine taste. {e} "

    try:
        # Gets the taste and usage recommendations
        pattern_usage = re.compile(
            r"\"usage\"\:\"[a-zA-Z\s\S]*\"\,\"alcoholPercentage\""
        )
        wine_usage = pattern_usage.findall(str(page_html))
        wine_usage = wine_usage[0][9:-21]
        # print(wine_taste_and_usage)
    except IndexError:
        wine_usage = "Failed to fetch wine usage recommendations. Regex 10 failure. "
    except Exception as e:
        wine_usage = f"Failed to fetch wine usage recommendations. {e} "

    # Combine taste and usage
    wine_taste_and_usage = wine_taste + " " + wine_usage

    # Create and fill the class with data and return it
    class_instance_wine = wine(
        wine_name,
        wine_name2,
        wine_location,
        wine_price,
        wine_productnr,
        wine_alc_procentage,
        wine_suger_content,
        wine_taste_and_usage,
    )
    return class_instance_wine


# Input list with wine classes
# Output creates output.txt file
def create_labels(class_list_wines):
    datum = date.today()

    new_file = open("output.txt", "w", encoding="utf-8")
    for wine in class_list_wines:
        new_file.write(wine.name + ", " + wine.name2 + "\n")
        new_file.write(wine.location + "\n")
        new_file.write(
            wine.price
            + "kr "
            + "sysnr:"
            + wine.productnr
            + "\t"
            + wine.alcohol_procentage
            + "% "
            + wine.suger_amount
            + "g/100ml \t"
            + str(datum)
            + "\n"
        )
        new_file.write(wine.taste_and_usage + "\n")
        """#Old print out style
        index = 0
        for text in wine.taste_and_usage:
            
            print(text, end='',file=new_file)
            if index%25 == 0 and index != 0 and 5 < (len(wine.taste_and_usage)-index):
                new_file.write("\n") 

            index +=1
        """

        new_file.write("\n\n\n\n")

    new_file.close()
    return


def create_output_docx(class_list_wines):
    datum = date.today()

    document = Document()
    for wine in class_list_wines:
        """
        document.add_paragraph(wine.name + ", " + wine.name2+ "\n")
        document.add_paragraph(wine.location + "\n")
        document.add_paragraph(wine.price +"kr "+ "sysnr:" + wine.productnr +"\t" + wine.alcohol_procentage + "% " + wine.suger_amount +"g/100ml \t" + str(datum) +"\n")
        document.add_paragraph(wine.taste_and_usage + "\n")
        document.add_paragraph()
        """
        document.add_paragraph(
            wine.name
            + ", "
            + wine.name2
            + "\n"
            + wine.location
            + "\n"
            + wine.price
            + "kr "
            + "sysnr:"
            + wine.productnr
            + "\t"
            + wine.alcohol_procentage
            + "% "
            + wine.suger_amount
            + "g/100ml \t"
            + str(datum)
            + "\n"
            + wine.taste_and_usage
            + "\n"
        )
        document.add_paragraph()

    document.save("output.docx")
    return


def main():

    filen_name = "input.txt"
    pages = load_url_file(filen_name)
    # print(pages) Test output of load_url_file

    pages = input_validation(pages)
    # print(pages) #check that the validation works

    site_data = []  # Variable list with entries being site data
    for index, url in enumerate(pages):
        site_data.append(dataurl_data_fetcher(url, index))

    class_list_wines = []
    for site in site_data:
        class_list_wines.append(extract_drink_data(site))

    try:  # remove for loops steppers.
        del site
    except:
        pass

    create_labels(class_list_wines)
    create_output_docx(class_list_wines)


if __name__ == "__main__":
    main()
