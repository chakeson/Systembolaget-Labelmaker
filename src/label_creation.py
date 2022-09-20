from datetime import date
from docx import Document


# Input list with wine classes
# Output creates output.txt file
def create_labels(class_list_wines):
    datum = date.today()

    new_file = open("output.txt", "w", encoding="utf-8")
    for wine in class_list_wines:
        new_file.write(wine.name + ", " + wine.name2 + "\n")
        new_file.write(wine.location + "\n")
        new_file.write(
            wine.price
            + "kr "
            + "sysnr:"
            + wine.productnr
            + "\t"
            + wine.alcohol_procentage
            + " "
            + wine.suger_amount
            + "\t"
            + str(datum)
            + "\n"
        )
        new_file.write(wine.taste_and_usage + "\n")
        """#Old print out style
        index = 0
        for text in wine.taste_and_usage:
            
            print(text, end='',file=new_file)
            if index%25 == 0 and index != 0 and 5 < (len(wine.taste_and_usage)-index):
                new_file.write("\n") 

            index +=1
        """

        new_file.write("\n\n\n\n")

    new_file.close()
    return


def create_output_docx(class_list_wines):
    datum = date.today()

    document = Document()
    for wine in class_list_wines:
        """
        document.add_paragraph(wine.name + ", " + wine.name2+ "\n")
        document.add_paragraph(wine.location + "\n")
        document.add_paragraph(wine.price +"kr "+ "sysnr:" + wine.productnr +"\t" + wine.alcohol_procentage + "% " + wine.suger_amount +"g/100ml \t" + str(datum) +"\n")
        document.add_paragraph(wine.taste_and_usage + "\n")
        document.add_paragraph()
        """
        document.add_paragraph(
            wine.name
            + ", "
            + wine.name2
            + "\n"
            + wine.location
            + "\n"
            + wine.price
            + "kr "
            + "sysnr:"
            + wine.productnr
            + "\t"
            + wine.alcohol_procentage
            + " "
            + wine.suger_amount
            + "\t"
            + str(datum)
            + "\n"
            + wine.taste_and_usage
            + "\n"
        )
        document.add_paragraph()

    document.save("output.docx")
    return